#!/usr/bin/env python3
"""Convert README.md to README.docx using python-docx."""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_code_block(doc: Document, code: str):
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.style = "No Spacing"
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        # Light grey background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F0F0F0")
        pPr.append(shd)


def apply_inline(paragraph, text: str):
    """Parse inline backtick code and bold within a text fragment."""
    # Split on backtick spans
    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            # Handle **bold**
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for bp in bold_parts:
                if bp.startswith('**') and bp.endswith('**'):
                    run = paragraph.add_run(bp[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(bp)


def parse_table(lines):
    """Parse a markdown table block, return list of row lists."""
    rows = []
    for line in lines:
        line = line.strip()
        if re.match(r'^\|[-:| ]+\|$', line):
            continue  # separator row
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    return rows


def build_docx(md_path: str, out_path: str):
    with open(md_path, encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = content.splitlines()
    i = 0
    in_code = False
    code_lines = []
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # ── Code block ─────────────────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_lines))
                doc.add_paragraph()  # spacer
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Markdown table ──────────────────────────────────────────────────────
        if line.startswith("|"):
            table_lines = [line]
            while i + 1 < len(lines) and lines[i + 1].startswith("|"):
                i += 1
                table_lines.append(lines[i])
            rows = parse_table(table_lines)
            if rows:
                ncols = len(rows[0])
                tbl = doc.add_table(rows=len(rows), cols=ncols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        apply_inline(p, cell_text)
                        if r_idx == 0:
                            p.runs[0].bold = True if p.runs else False
                            set_cell_bg(cell, "D9E1F2")
                doc.add_paragraph()
            i += 1
            continue

        # ── Horizontal rule ─────────────────────────────────────────────────────
        if re.match(r'^---+$', line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Headings ────────────────────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            heading_map = {1: "Heading 1", 2: "Heading 2",
                           3: "Heading 3", 4: "Heading 4"}
            p = doc.add_heading(level=level)
            p.style = doc.styles[heading_map[level]]
            apply_inline(p, text)
            i += 1
            continue

        # ── Bullet list ─────────────────────────────────────────────────────────
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            indent = len(m.group(1)) // 2
            text   = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
            apply_inline(p, text)
            i += 1
            continue

        # ── Numbered list ───────────────────────────────────────────────────────
        m = re.match(r'^\d+\.\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style="List Number")
            apply_inline(p, m.group(1))
            i += 1
            continue

        # ── Empty line ──────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Normal paragraph ────────────────────────────────────────────────────
        p = doc.add_paragraph()
        apply_inline(p, line)
        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    import os
    here = os.path.dirname(__file__)
    build_docx(
        md_path=os.path.join(here, "README.md"),
        out_path=os.path.join(here, "README.docx"),
    )
